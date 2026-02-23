import io
import json
import os
import requests
from .tasks import ask_learnflow_ai
from django.http import JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.contrib import messages
from django.http import Http404, HttpResponse
from django.template.loader import get_template
from django.db import transaction
from django.utils import timezone
from django.core.exceptions import ObjectDoesNotExist
from docx import Document
import requests
from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import render, redirect # Ensure this is imported at the top
from .models import Quiz, Question, Choice # Ensure these are imported
from .forms import QuizForm # Ensure this is imported
# Import the necessary libraries for PDF generation
from xhtml2pdf import pisa
# Import models and forms from both aiapp and video apps
from .models import Quiz, Question, Choice, StudentAnswer, Attempt
from .forms import QuizForm
from video.models import Video
from video.forms import VideoForm
from book.models import Book
from django.contrib.auth.decorators import login_required
from urllib.parse import urljoin





# from django.views.decorators.csrf import csrf_exempt

# NOTE: For this to work with your AI models, you'll need to import
# and use them here. This is a placeholder for the actual AI logic.
 # adjust if Quiz is in another app

@login_required
def home(request):
    """
    Renders the AI-powered home page with personalized context.
    """
    return render(request, 'aiapp/home.html', {
        "user": request.user,
        "username": request.user.get_full_name() or request.user.username,
        "ai_context": "Praise AI is here to empower educators and learners across Africa. Ask anything!",
         "show_ads": True
    })


def ask_ai_view(request):
    query = request.POST.get("query")
    task = ask_learnflow_ai.delay(query)
    return JsonResponse({"task_id": task.id})

def render_to_pdf(template_src, context_dict={}):
    """
    Renders a Django template to a PDF file.
    """
    try:
        template = get_template(template_src)
        html = template.render(context_dict)
        result = io.BytesIO()
        pdf = pisa.pisaDocument(io.BytesIO(html.encode("UTF-8")), result)
        if not pdf.err:
            return HttpResponse(result.getvalue(), content_type='application/pdf')
        else:
            print("PDF generation error:", pdf.err)
            print("Rendered HTML:", html)
    except Exception as e:
        print("Exception during PDF generation:", str(e))
    return None



@login_required
def quiz_list(request):
    """
    Renders a list of all quizzes for students to view.
    """
    quizzes = Quiz.objects.order_by('-created_at')
    return render(request, 'aiapp/quiz_list.html', {
        'quizzes': quizzes,
        'show_ads': True
        })

@login_required
def quiz_detail(request, quiz_id):
    """
    Displays the details of a specific quiz.
    """
    quiz = get_object_or_404(Quiz, id=quiz_id)
    return render(request, 'aiapp/quiz_detail.html', {
        'quiz': quiz,
        'show_ads': True
    })

@login_required
@transaction.atomic
def quiz_attempt(request, quiz_id):
    """
    Handles the student's attempt at a quiz.
    """
    quiz = get_object_or_404(Quiz, id=quiz_id)
    questions = quiz.questions.all()

    if request.method == 'POST':
        score = 0
        total_questions = questions.count()

        new_attempt = Attempt.objects.create(
            user=request.user,
            quiz=quiz,
            score=0,
            total_questions=total_questions
        )

        for question in questions:
            is_correct = False
            selected_choice = None
            submitted_answer_text = None

            if question.question_type == 'MC':
                submitted_choice_id = request.POST.get(f'question_{question.id}', None)
                if submitted_choice_id:
                    try:
                        selected_choice = Choice.objects.get(pk=int(submitted_choice_id))
                        is_correct = selected_choice.is_correct
                    except (ValueError, ObjectDoesNotExist):
                        is_correct = False
            
            elif question.question_type == 'SA':
                submitted_answer_text = request.POST.get(f'question_{question.id}', '')
                if submitted_answer_text.strip():
                    is_correct = (submitted_answer_text.strip().lower() == question.correct_answer_text.strip().lower())
                else:
                    is_correct = False
            
            StudentAnswer.objects.create(
                student=request.user,
                question=question,
                selected_choice=selected_choice,
                text_answer=submitted_answer_text,
                is_correct=is_correct,
                attempt=new_attempt
            )
            
            if is_correct:
                score += 1
        
        new_attempt.score = score
        new_attempt.save()

        return redirect('aiapp:quiz_results', attempt_id=new_attempt.id)

    return render(request, 'aiapp/quiz_attempt.html', {
        'quiz': quiz, 
        'questions': questions,
        'show_ads': True
        })

@login_required
def quiz_results(request, attempt_id):
    """
    Displays the results of a specific quiz attempt.
    """
    try:
        attempt = get_object_or_404(Attempt, pk=attempt_id, user=request.user)
        passed = attempt.score >= (attempt.total_questions * 0.5)

        # Calculate the percentage and the SVG stroke-dashoffset here
        if attempt.total_questions > 0:
            score_percentage = attempt.score / attempt.total_questions
            score_offset = 339.29 * (1 - score_percentage)
        else:
            score_percentage = 0
            score_offset = 339.29 # Full circle offset if no questions

        return render(request, 'aiapp/quiz_results.html', {
            'quiz': attempt.quiz,
            'score': attempt.score,
            'total_questions': attempt.total_questions,
            'attempt_id': attempt.id,
            'passed': passed,
            'score_percentage': score_percentage,
            'score_offset': score_offset,
            'show_ads': True
        })
    except Exception as e:
        print("Error rendering quiz results:", str(e))
        return HttpResponse("Error displaying results", status=500)

@login_required
def quiz_review(request, attempt_id):
    """
    Displays a detailed review of a specific quiz attempt.
    """
    attempt = get_object_or_404(Attempt, pk=attempt_id, user=request.user)
    student_answers = StudentAnswer.objects.filter(attempt=attempt).select_related('question', 'selected_choice')
    
    questions_and_answers = []
    for q in attempt.quiz.questions.all():
        try:
            student_answer = student_answers.get(question=q)
        except ObjectDoesNotExist:
            student_answer = None
        
        questions_and_answers.append({
            'question': q,
            'student_answer': student_answer,
        })

    context = {
        'attempt': attempt,
        'questions_and_answers': questions_and_answers,
        'show_ads': True
    }
    return render(request, 'aiapp/quiz_review.html', context)


# @login_required
# @transaction.atomic
# def create_quiz(request):
#     """
#     Handles the creation of a new quiz and its questions.
#     """
#     if request.method == 'POST':
#         quiz_form = QuizForm(request.POST)
        
#         if quiz_form.is_valid():
#             quiz = quiz_form.save(commit=False)
#             quiz.teacher = request.user
#             quiz.upload_code = quiz_form.cleaned_data.get('upload_code')
#             quiz.save()
            
            # questions_data = request.POST.get('questions_data')
            # if questions_data:
            #     try:
            #         questions_list = json.loads(questions_data)
            #         for q_data in questions_list:
            #             question_text = q_data.get('text')
            #             if not question_text:
            #                 continue

            #             question_type = q_data.get('question_type', 'MC')
                        
            #             question_instance = Question.objects.create(
            #                 quiz=quiz,
            #                 text=question_text,
            #                 question_type=question_type,
            #             )

                        # if question_type == 'MC':
                        #     choices_data = q_data.get('choices', [])
                        #     for c_data in choices_data:
                        #         if c_data.get('text'):
                        #             Choice.objects.create(
                        #                 question=question_instance,
                        #                 text=c_data['text'],
                        #                 is_correct=c_data.get('isCorrect', False)
                        #             )
                        # elif question_type == 'SA':
                        #     correct_answer_text = q_data.get('correct_answer_text')
                        #     if correct_answer_text:
                        #         question_instance.correct_answer_text = correct_answer_text
                        #         question_instance.save()

    #                 messages.success(request, f'"{quiz.title}" has been created successfully!')
    #                 return redirect('aiapp:teacher_quiz_dashboard')

    #             except (json.JSONDecodeError, KeyError) as e:
    #                 print(f"JSON processing error: {e}")
    #                 messages.error(request, "There was an error processing the quiz questions. Please check the data and try again.")
    #                 raise
    #     else:
    #         messages.error(request, "Please correct the form errors below.")
    #         print("Form errors:", quiz_form.errors)
    # else:
    #     quiz_form = QuizForm()
        
    # return render(request, 'aiapp/create_quiz.html', {
    #     'quiz_form': quiz_form,
    #     'show_ads': True
    #     })
       

@login_required
def teacher_quiz_dashboard(request):
    """
    Displays a dashboard of quizzes created by the current user.
    """
    user_quizzes = Quiz.objects.filter(teacher=request.user).order_by('-created_at')
    return render(request, 'aiapp/teacher_quiz_dashboard.html', {
        'user_quizzes': user_quizzes,
        'show_ads': True
        })

@login_required
@transaction.atomic
def edit_quiz(request, quiz_id):
    """
    Allows a teacher to edit an existing quiz.
    """
    quiz = get_object_or_404(Quiz, pk=quiz_id)
    
    if quiz.teacher != request.user:
        raise Http404

    if request.method == 'POST':
        quiz_form = QuizForm(request.POST, instance=quiz)
        if quiz_form.is_valid():
            quiz = quiz_form.save()
            
            questions_data = request.POST.get('questions_data')
            if questions_data:
                try:
                    questions_list = json.loads(questions_data)
                    
                    existing_question_ids = set(quiz.questions.values_list('id', flat=True))
                    questions_to_keep = set()

                    for q_data in questions_list:
                        question_id = q_data.get('id')
                        question_text = q_data.get('text')
                        if not question_text:
                            continue
                        
                        question_type = q_data.get('question_type', 'MC')
                        
                        if question_id and str(question_id).isdigit():
                            question_instance = Question.objects.get(pk=question_id, quiz=quiz)
                            question_instance.text = question_text
                            question_instance.question_type = question_type
                            
                            questions_to_keep.add(int(question_id))
                        else:
                            question_instance = Question.objects.create(
                                quiz=quiz,
                                text=question_text,
                                question_type=question_type,
                            )
                        
                        if question_type == 'MC':
                            question_instance.choice_set.all().delete()
                            choices_data = q_data.get('choices', [])
                            for c_data in choices_data:
                                if c_data.get('text'):
                                    Choice.objects.create(
                                        question=question_instance,
                                        text=c_data['text'],
                                        is_correct=c_data.get('isCorrect', False)
                                    )
                            question_instance.correct_answer_text = None
                            question_instance.save()
                        elif question_type == 'SA':
                            correct_answer_text = q_data.get('correct_answer_text')
                            if correct_answer_text:
                                question_instance.correct_answer_text = correct_answer_text
                            else:
                                question_instance.correct_answer_text = ""
                            question_instance.save()
                            question_instance.choice_set.all().delete()

                    questions_to_delete_ids = existing_question_ids - questions_to_keep
                    Question.objects.filter(id__in=questions_to_delete_ids, quiz=quiz).delete()
                                        
                    messages.success(request, f'"{quiz.title}" has been updated successfully!')
                    return redirect('aiapp:teacher_quiz_dashboard')

                except (json.JSONDecodeError, KeyError, ObjectDoesNotExist) as e:
                    print(f"JSON processing error: {e}")
                    messages.error(request, "There was an error processing the quiz questions. Please check the data and try again.")
                    raise
        else:
            messages.error(request, "Please correct the form errors below.")
    else:
        quiz_form = QuizForm(instance=quiz)

    return render(request, 'aiapp/edit_quiz.html', {'quiz_form': quiz_form,
                                                    'quiz': quiz,
                                                    'show_ads': True
                                                    })

@login_required
def delete_quiz(request, quiz_id):
    """
    Allows a teacher to delete a quiz.
    """
    quiz = get_object_or_404(Quiz, pk=quiz_id)
    
    if quiz.teacher != request.user:
        raise Http404
        
    if request.method == 'POST':
        quiz.delete()
        messages.success(request, f'"{quiz.title}" has been deleted successfully!')
        return redirect('aiapp:teacher_quiz_dashboard')

    return render(request, 'aiapp/delete_quiz_confirm.html', {'quiz': quiz, 'show_ads': True})

@login_required
def user_profile(request, user_id):
    """
    Displays the user profile page.
    """
    user_to_display = get_object_or_404(User, pk=user_id)
    return render(request, 'aiapp/profile_detail.html', {'profile_user': user_to_display, 'show_ads': True})

# Video-related views from the video app
@login_required
def video_list(request):
    """
    Renders a list of all available videos.
    """
    videos = Video.objects.all().order_by('-created_at')
    return render(request, 'video/video_list.html', {'videos': videos, 'show_ads': True})

@login_required
def video_detail(request, video_id):
    """
    Renders the detail page for a single video.
    """
    video = get_object_or_404(Video, pk=video_id)
    return render(request, 'video/video_detail.html', {'video': video, 'show_ads': True})

@login_required
def create_video(request):
    """
    Allows a teacher to upload a new video.
    """
    if request.method == 'POST':
        form = VideoForm(request.POST)
        if form.is_valid():
            video = form.save(commit=False)
            video.teacher = request.user
            video.save()
            form.save_m2m()
            messages.success(request, f'"{video.title}" has been uploaded successfully!')
            return redirect('video:teacher_dashboard')
    else:
        form = VideoForm()
    return render(request, 'video/video_upload.html', {'form': form, 'show_ads': True})

@login_required
def teacher_dashboard(request):
    """
    Displays a dashboard of videos uploaded by the current user.
    """
    user_videos = Video.objects.filter(teacher=request.user).order_by('-created_at')
    return render(request, 'video/teacher_dashboard.html', {'user_videos': user_videos, 'show_ads': True})

@login_required
def edit_video(request, video_id):
    """
    Allows a teacher to edit an existing video.
    """
    video = get_object_or_404(Video, pk=video_id)
    
    if video.teacher != request.user:
        raise Http404
        
    if request.method == 'POST':
        form = VideoForm(request.POST, instance=video)
        if form.is_valid():
            form.save()
            messages.success(request, f'"{video.title}" has been updated successfully!')
            return redirect('video:teacher_dashboard')
    else:
        form = VideoForm(instance=video)

    return render(request, 'video/video_edit.html', {'form': form, 'video': video, 'show_ads': True})

@login_required
def delete_video(request, video_id):
    """
    Allows a teacher to delete a video.
    """
    video = get_object_or_404(Video, pk=video_id)
    
    if video.teacher != request.user:
        raise Http404
        
    if request.method == 'POST':
        video.delete()
        messages.success(request, f'"{video.title}" has been deleted successfully!')
        return redirect('video:teacher_dashboard')

    return render(request, 'video/video_delete_confirm.html', {'video': video, 'show_ads': True})

@login_required
def quiz_report_pdf_for_quiz(request, quiz_id):
    """Generates a PDF report for a specific quiz, including all student attempts."""
    quiz = get_object_or_404(Quiz, pk=quiz_id)
    if request.user != quiz.teacher:
        raise Http404

    attempts = Attempt.objects.filter(quiz=quiz)
    enriched_attempts = []

    for a in attempts:
        percentage = round((a.score / a.total_questions) * 100) if a.total_questions else 0
        incorrect = a.total_questions - a.score
        enriched_attempts.append({
            'user': a.user,
            'score': a.score,
            'total_questions': a.total_questions,
            'percentage': percentage,
            'incorrect_answers': incorrect,
        })

    context = {
        'quiz': quiz,
        'attempts': enriched_attempts,
        'report_date': timezone.now(),
        'request_user': request.user,
        'show_ads': True
    }

    try:
        pdf = render_to_pdf('aiapp/quiz_report_pdf.html', context)
        if pdf:
            filename = f"{quiz.title.replace(' ', '_')}_report.pdf"
            response = HttpResponse(pdf, content_type='application/pdf')
            response['Content-Disposition'] = f"attachment; filename='{filename}'"
            return response
    except Exception as e:
        print("PDF generation error:", str(e))

    return HttpResponse("Error generating PDF", status=500)

def build_quiz_attempt_context(attempt, request_user):
    student_answers = StudentAnswer.objects.filter(attempt=attempt).select_related('question', 'selected_choice')
    score = attempt.score
    total_questions = attempt.total_questions
    percentage = round((score / total_questions) * 100) if total_questions else 0
    incorrect_answers = total_questions - score

    enriched_answers = []
    for ans in student_answers:
        correct_choice = ans.question.choices.filter(is_correct=True).first()
        user_answer = ans.selected_choice.text if ans.selected_choice else ans.text_answer
        correct_answer = correct_choice.text if correct_choice else "N/A"
        is_correct = user_answer == correct_answer
        feedback = "✅ Well done!" if is_correct else f"❌ Almost there — the correct answer was '{correct_answer}'. Keep going!"

        enriched_answers.append({
            'question_text': ans.question.text,
            'user_answer': user_answer,
            'correct_answer': correct_answer,
            'is_correct': is_correct,
            'feedback': feedback,
        })

    return {
        'quiz': attempt.quiz,
        'attempt': attempt,
        'report_date': timezone.now(),
        'request_user': request_user,
        'score': score,
        'total_questions': total_questions,
        'percentage': percentage,
        'incorrect_answers': incorrect_answers,
        'user': attempt.user,
        'today': timezone.now(),
        'answers': enriched_answers,
        'show_ads': True
    }    

@login_required
def quiz_report_pdf_for_attempt(request, attempt_id):
    """Generates a PDF report for a specific quiz attempt, including correct answers and feedback."""
    attempt = get_object_or_404(Attempt, pk=attempt_id)
    if request.user != attempt.user and request.user != attempt.quiz.teacher:
        raise Http404

    student_answers = StudentAnswer.objects.filter(attempt=attempt).select_related('question', 'selected_choice')
    score = attempt.score
    total_questions = attempt.total_questions
    percentage = round((score / total_questions) * 100) if total_questions else 0
    incorrect_answers = total_questions - score

    enriched_answers = []
    for ans in student_answers:
        user_answer = ans.selected_choice.text if ans.selected_choice else ans.text_answer
        
        # 1. New variable for correct answer text
        correct_answer_text_for_display = ""
        
        # 2. Conditional logic for correct_answer_text_for_display
        if ans.question.question_type == 'MC':
            correct_choice = ans.question.choices.filter(is_correct=True).first()
            correct_answer_text_for_display = correct_choice.text if correct_choice else "N/A (No correct choice defined)"
        elif ans.question.question_type == 'SA':
            correct_answer_text_for_display = ans.question.correct_answer_text if ans.question.correct_answer_text else "N/A (No correct answer defined)"
        else:
            correct_answer_text_for_display = "N/A (Unknown question type)"
        
        # 3. Direct use of ans.is_correct
        is_correct = ans.is_correct 

        # 4. Updated feedback message
        feedback = (
            "✅ Well done!" if is_correct else
            f"❌ Almost there — the correct answer was '{correct_answer_text_for_display}'. Keep going!"
        )

        # 5. Update in the enriched_answers dictionary
        enriched_answers.append({
            'question_text': ans.question.text,
            'user_answer': user_answer,
            'correct_answer': correct_answer_text_for_display, # Use the correctly determined text
            'is_correct': is_correct, # Add this flag for potential template use
            'feedback': feedback,
        })

    context = {
        'quiz': attempt.quiz,
        'attempt': attempt,
        'report_date': timezone.now(),
        'request_user': request.user,
        'score': score,
        'total_questions': total_questions,
        'percentage': percentage,
        'incorrect_answers': incorrect_answers,
        'user': attempt.user,
        'today': timezone.now(),
        'answers': enriched_answers,
    }

    try:
        pdf = render_to_pdf('aiapp/quiz_report_pdf.html', context)
        if pdf:
            filename = f"{attempt.quiz.title.replace(' ', '_')}_report_attempt_{attempt.id}.pdf"
            response = HttpResponse(pdf, content_type='application/pdf')
            response['Content-Disposition'] = f"attachment; filename='{filename}'"
            return response
    except Exception as e:
        print("PDF generation error:", str(e))

    return HttpResponse("Error generating PDF", status=500)


def build_quiz_attempt_context(attempt, request_user):
    student_answers = StudentAnswer.objects.filter(attempt=attempt).select_related('question', 'selected_choice')
    score = attempt.score
    total_questions = attempt.total_questions
    percentage = round((score / total_questions) * 100) if total_questions else 0
    incorrect_answers = total_questions - score

    enriched_answers = []
    for ans in student_answers:
        user_answer = ans.selected_choice.text if ans.selected_choice else ans.text_answer
        
        # 1. New variable for correct answer text
        correct_answer_text_for_display = ""
        
        # 2. Conditional logic for correct_answer_text_for_display
        if ans.question.question_type == 'MC':
            correct_choice = ans.question.choices.filter(is_correct=True).first()
            correct_answer_text_for_display = correct_choice.text if correct_choice else "N/A (No correct choice defined)"
        elif ans.question.question_type == 'SA':
            correct_answer_text_for_display = ans.question.correct_answer_text if ans.question.correct_answer_text else "N/A (No correct answer defined)"
        else:
            correct_answer_text_for_display = "N/A (Unknown question type)"
        
        # 3. Direct use of ans.is_correct
        is_correct = ans.is_correct 

        # 4. Updated feedback message
        feedback = (
            "✅ Well done!" if is_correct else
            f"❌ Almost there — the correct answer was '{correct_answer_text_for_display}'. Keep going!"
        )

        # 5. Update in the enriched_answers dictionary
        enriched_answers.append({
            'question_text': ans.question.text,
            'user_answer': user_answer,
            'correct_answer': correct_answer_text_for_display,
            'is_correct': is_correct,
            'feedback': feedback,
        })

    return {
        'quiz': attempt.quiz,
        'attempt': attempt,
        'report_date': timezone.now(),
        'request_user': request_user,
        'score': score,
        'total_questions': total_questions,
        'percentage': percentage,
        'incorrect_answers': incorrect_answers,
        'user': attempt.user,
        'today': timezone.now(),
        'answers': enriched_answers,
    }    
# ----------------------------------------------------------------------

@login_required
def quiz_report_word_for_attempt(request, attempt_id):
    attempt = get_object_or_404(Attempt, pk=attempt_id)
    if request.user != attempt.user and request.user != attempt.quiz.teacher:
        raise Http404

    # This call now returns the context with correctly determined correct_answer and is_correct flags
    context = build_quiz_attempt_context(attempt, request.user)
    
    doc = Document()
    doc.add_heading(f'Quiz Report - {context["quiz"].title}', 0)
    doc.add_paragraph(f'Generated for {context["user"].username} on {context["today"].strftime("%B %d, %Y")}')

    doc.add_heading('Summary', level=1)
    doc.add_paragraph(f'Score: {context["score"]} / {context["total_questions"]}')
    doc.add_paragraph(f'Correct: {context["score"]}')
    doc.add_paragraph(f'Incorrect: {context["incorrect_answers"]}')
    doc.add_paragraph(f'Percentage: {context["percentage"]}%')

    doc.add_heading('Questions and Answers', level=1)
    for i, ans in enumerate(context['answers'], start=1):
        doc.add_paragraph(f'Question {i}: {ans["question_text"]}', style='List Number')
        doc.add_paragraph(f'Your Answer: {ans["user_answer"]}')
        # This key now contains the correct text for MC or SA questions
        doc.add_paragraph(f'Correct Answer: {ans["correct_answer"]}')
        # This key contains the correct feedback string
        doc.add_paragraph(f'Feedback: {ans["feedback"]}')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"{context['quiz'].title.replace(' ', '_')}_attempt_{attempt.id}.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    doc.save(response)
    return response

def why_learnflow_ai(request):
    return render(request, 'aiapp/why_learnflow_ai.html')

def ai_quiz_generator(request):
    return render(request, 'aiapp/ai_quiz_generator.html')


def sitemap_view(request):
    # Detect host and set base_url accordingly
    host = request.get_host()
    if "onrender.com" in host:
        base_url = "https://learnflow-ai-0fdz.onrender.com"
    elif "koyeb.app" in host:
        base_url = "https://artificial-shirlee-learnflow-8ec0e7a0.koyeb.app"
    else:
        base_url = request.build_absolute_uri("/")[:-1]  # fallback to current host

    urls = [
        urljoin(base_url, "/"),
        urljoin(base_url, "/aiapp/"),
        urljoin(base_url, "/video/"),
        urljoin(base_url, "/book/"),
    ]

    try:
        for q in Quiz.objects.all():
            slug_or_pk = getattr(q, 'slug', None) or str(q.pk)
            urls.append(urljoin(base_url, f"/quiz/{slug_or_pk}/"))
    except Exception:
        urls.append(urljoin(base_url, "/quiz/"))

    try:
        for b in Book.objects.all():
            slug_or_pk = getattr(b, 'slug', None) or str(b.pk)
            urls.append(urljoin(base_url, f"/books/{slug_or_pk}/"))
    except Exception:
        urls.append(urljoin(base_url, "/books/"))

    try:
        for v in Video.objects.all():
            slug_or_pk = getattr(v, 'slug', None) or str(v.pk)
            urls.append(urljoin(base_url, f"/video/{slug_or_pk}/"))
    except Exception:
        urls.append(urljoin(base_url, "/video/"))

    xml = '<?xml version="1.0" encoding="UTF-8"?>\n'
    xml += '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">\n'
    for url in urls:
        xml += f'  <url><loc>{url}</loc></url>\n'
    xml += '</urlset>'

    return HttpResponse(xml, content_type='application/xml')

@login_required
def retake_quiz(request, quiz_id):
    quiz = get_object_or_404(Quiz, pk=quiz_id)
    return redirect('aiapp:quiz_attempt', quiz_id=quiz.id)



# Replace with your actual Botlhale token
BOTLHALE_API_TOKEN = "Bearer YOUR_BOTLHALE_TOKEN"
BOTLHALE_TTS_URL = "https://api.botlhale.xyz/tts"

@csrf_exempt
def tts_proxy(request):
    if request.method != "POST":
        return JsonResponse({"error": "Only POST requests are allowed."}, status=405)

    try:
        data = json.loads(request.body)
        text = data.get("text")
        language_code = data.get("language_code")

        if not text or not language_code:
            return JsonResponse({"error": "Missing 'text' or 'language_code'."}, status=400)

        payload = {
            "text": text,
            "language_code": language_code
        }

        headers = {
            "Authorization": BOTLHALE_API_TOKEN
        }

        response = requests.post(BOTLHALE_TTS_URL, headers=headers, data=payload)

        if response.status_code != 200:
            return JsonResponse({"error": "Botlhale API error", "details": response.text}, status=response.status_code)

        return JsonResponse(response.json())

    except Exception as e:
        return JsonResponse({"error": "Server error", "details": str(e)}, status=500)




@login_required
def create_quiz(request):
    """
    Handles the creation of a new quiz, including its questions and choices,
    from the POST request containing QuizForm data and questions_json.
    """
    if request.method == 'POST':
        # The upload_code is reconstructed by JS into the hidden 'id_upload_code' input 
        # which corresponds to the QuizForm field 'upload_code'.
        quiz_form = QuizForm(request.POST)
        
        # The form submission logic relies on a hidden 'questions_json' field
        questions_json = request.POST.get('questions_json')
        
        if quiz_form.is_valid():
            
            if not questions_json:
                messages.error(request, "Please add at least one question to the quiz.")
                # Assumes the template is correctly named based on previous context
                return render(request, 'aiapp/create_quiz.html', {'quiz_form': quiz_form}) 

            try:
                questions_data = json.loads(questions_json)
            except (json.JSONDecodeError, TypeError):
                messages.error(request, "Invalid question data submitted.")
                return render(request, 'aiapp/create_quiz.html', {'quiz_form': quiz_form})
            
            if not questions_data:
                messages.error(request, "Please add at least one question to the quiz.")
                return render(request, 'aiapp/create_quiz.html', {'quiz_form': quiz_form})


            # 1. Save the Quiz object
            with transaction.atomic():
                # The logged-in user is the teacher
                quiz = quiz_form.save(commit=False)
                # Assuming request.user is the teacher as per @login_required
                quiz.teacher = request.user 
                quiz.save()
                
                # 2. Process Questions and Choices from JSON
                for q_data in questions_data:
                    # Create the Question
                    question = Question.objects.create(
                        quiz=quiz,
                        text=q_data['text'],
                        # FIX APPLIED HERE: Changed 'type' to 'question_type' 
                        # to match the field name in models.py
                        question_type=q_data['type'] 
                    )

                    if q_data['type'] == 'MC': # Multiple Choice
                        for c_data in q_data['choices']:
                            Choice.objects.create(
                                question=question,
                                text=c_data['text'],
                                is_correct=c_data['is_correct']
                            )
                    elif q_data['type'] == 'SA': # Single Answer
                        # For Single Answer, the correct text is stored as a single Choice 
                        # linked to the question, with is_correct=True.
                        Choice.objects.create(
                            question=question,
                            text=q_data['correct_answer'],
                            is_correct=True 
                        )

                messages.success(request, f'Quiz "{quiz.title}" created successfully!')
                return redirect('aiapp:teacher_quiz_dashboard') # Redirect to the teacher dashboard

        else:
            # Form is invalid (e.g., title missing, upload_code validation failed)
            messages.error(request, "There was an error with your quiz details. Please check the form. Note: Upload Access Code must be valid and 5 digits.")
    
    else:
        # GET request: render the empty form
        quiz_form = QuizForm()
        
    return render(request, 'aiapp/create_quiz.html', {'quiz_form': quiz_form})

# --- Utility function to clean history format ---
def clean_contents(messages):
    """
    Ensures messages conform to the Gemini API format:
    1. Standardizes role: 'Ai' -> 'model'
    2. Standardizes content: 'text' property is moved into 'parts' array.
    """
    cleaned = []
    for msg in messages:
        # Standardize role: 'Ai' -> 'model'
        role = msg.get("role", "").lower()
        if role == "ai":
            role = "model"
        
        # Standardize content format: 'text' -> 'parts'
        # This handles the legacy format from localStorage
        if "text" in msg and not msg.get("parts"):
            cleaned.append({
                "role": role,
                "parts": [{"text": msg["text"]}]
            })
        else:
            # Assume it's already in the correct Gemini format
            cleaned.append(msg)
    return cleaned
# ---------------------------------------------


@csrf_exempt
def gemini_proxy(request):
    if request.method != "POST":
        return JsonResponse({"error": "Only POST allowed"}, status=405)

    try:
        # Request body handling
        body = json.loads(request.body.decode("utf-8"))

        # 1. Setup API Key and URL 
        # We prioritize the OS environment key (for Koyeb), but fall back to the Canvas global key
        api_key = os.environ.get("GEMINI_API_KEY")
        if not api_key:
            # Fallback for local testing/Canvas environment where the key is auto-injected
            api_key = globals().get('__api_key', '')
            if not api_key:
                 # This check should theoretically never be reached in the Canvas environment
                 return JsonResponse({"error": "Missing GEMINI_API_KEY in environment or globals."}, status=500)


        model = "gemini-2.5-flash"
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"

        # 2. Extract and Clean Contents (Chat History)
        contents = body.get("contents")
        if not contents:
            # Fallback for testing
            contents = [{"role": "user", "parts": [{"text": "Hello Gemini"}]}]

        # CRITICAL FIX: Clean the history structure before sending
        contents = clean_contents(contents)

        # 3. Define System Instruction Text
        system_instruction_text = (
            "You are Praise AI, an educational partner developed by Kintu Peter, "
            "CEO of Mwene Groups of Companies. Always provide accurate, empathetic, and concise answers."
        )
        
        # CRITICAL FIX: Format the System Instruction as a Content object for the REST API
        system_instruction_content = {
            "role": "system",
            "parts": [{"text": system_instruction_text}]
        }

        # 4. Construct Generation Config (Type Casting Fixes)
        config = body.get("config") or {}
        generation_config = {
            # Ensure temperature is explicitly cast to a float
            "temperature": float(config.get("temperature", 0.7)), 
            # Ensure maxOutputTokens is explicitly cast to an integer
            "maxOutputTokens": int(config.get("maxOutputTokens", 1024)),
        }

        # 5. Construct Final Payload
        payload = {
            "contents": contents,
            # PASS THE CORRECTLY STRUCTURED CONTENT OBJECT
            "systemInstruction": system_instruction_content, 
            "generationConfig": generation_config,
        }

        # Debug log: Log the final outbound payload before sending
        print("Outbound Gemini payload:", json.dumps(payload, indent=2))

        # 6. Make the API Request
        resp = requests.post(url, json=payload)
        
        # 7. Error Handling
        if resp.status_code != 200:
            # CRITICAL LOG: This will show the exact reason for the 400 error.
            print("Gemini API Error Details:", resp.text)
            return JsonResponse(
                {"error": f"Gemini API error {resp.status_code}", "details": resp.text},
                status=resp.status_code,
            )

        # 8. Success Response Handling
        data = resp.json()
        text = ""
        if "candidates" in data and data["candidates"]:
            # Extracting text from parts, accommodating multiple parts if they exist
            parts = data["candidates"][0].get("content", {}).get("parts", [])
            text = " ".join(p.get("text", "") for p in parts if "text" in p)

        return JsonResponse({"text": text, "raw": data})

    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)